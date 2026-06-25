"""Cleaned analysis script converted from Jupyter notebook."""

import numpy as np
import pandas as pd
from sklearn.model_selection import GridSearchCV
from xgboost import XGBRegressor
from sklearn.model_selection import train_test_split
import matplotlib.pyplot as plt
from econml.dml import LinearDML
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import KFold
from sklearn.metrics import mean_squared_error
from sklearn.impute import SimpleImputer
import networkx as nx
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import seaborn as sns
from dowhy import CausalModel
from matplotlib.patches import FancyBboxPatch


# Load dataset
data = pd.read_csv('combined_aar_model.csv')

# Split the dataset into outcome, treatment, and covariates
Y = data.iloc[:, 0].values  # Outcome variable
T = data.iloc[:, 1].values  # Treatment variable
X = data.iloc[:, 2:].values  # Covariates

# Split the data into training and testing sets
X_train, X_test, Y_train, Y_test, T_train, T_test = train_test_split(X, Y, T, test_size=0.3, random_state=42)

# Hyperparameter optimization for XGBoost using GridSearchCV
xgb_model = XGBRegressor(objective='reg:squarederror')
param_grid = {
    'n_estimators': [50, 100, 200, 300],
    'learning_rate': [0.01, 0.1, 0.2, 0.3],
    'max_depth': [3, 5, 7, 9],
    'subsample': [0.8, 1.0],
    'colsample_bytree': [0.8, 1.0]
}

grid_search = GridSearchCV(estimator=xgb_model, param_grid=param_grid, cv=3, n_jobs=-1, verbose=2)
grid_search.fit(X_train, Y_train)

# Best parameters 
best_params = grid_search.best_params_

# Print best parameters
print(best_params)

# Create XGBoost models with the best parameters for DML
model_y = XGBRegressor(**best_params)
model_t = XGBRegressor(**best_params)

# Linear DML estimator
est = LinearDML(model_y=model_y, model_t=model_t, linear_first_stages=False, cv=3)

# Fit model
est.fit(Y_train, T_train, X=X_train)

# Estimate the marginal effects of the treatment by each covariate
marginal_effects = est.marginal_effect(T_test, X=X_test)

# MANUAL DOUBLE MACHINE LEARNING
# No econml / no doubleml

# IMPORT LIBRARIES

# LOAD DATA

df = pd.read_csv("combined.csv")

# KEEP VISIT = 4 ONLY

df = df[df["visit"] == 4].copy()

print(df.shape)

# REMOVE HIV PROXY VARIABLES
# + INTERACTION TERMS

remove_keywords = [

    "cvl",
    "vload2",
    "esthivdate",
    "lg10cvl",
    "cvlfrhaart"

]

pattern = "|".join(remove_keywords)

df = df.loc[
    :,
    ~df.columns.str.contains(
        pattern,
        case=False,
        regex=True
    )
]

print(df.shape)

# DEFINE OUTCOMES

outcomes = [

    "aar",
    "eeaa",
    "peaa",
    "geaa",
    "dnamtladjage"

]

# DEFINE TREATMENTS

treatments = [

    "Precipitation",
    "Temperature",
    "SPWPM2.5",
    "SBMPM2.5",
    "CO",
    "COAQI",
    "NO2",
    "NO2AQI",
    "Ozone",
    "OzoneAQI",
    "SO2",
    "SO2AQI"

]

# DEFINE COVARIATES

covariates = [

    "white",
    "educbas",
    "bmi",
    "cum_pkyear",

    "lncd3",
    "lncd4",
    "lncd8",

    "lnnaivecd4",
    "lnnaivecd8",

    "lnsencd4",
    "lnsencd8",

    "lnactcd4",
    "lnactcd8",

    "lnabs_cd42",
    "lnabs_cd82",

    "lnabsnaivecd4",
    "lnabsnaivecd8",

    "lnabssencd4",
    "lnabssencd8",

    "lnabsactcd4",
    "lnabsactcd8"

]

# KEEP EXISTING VARIABLES

outcomes = [
    x for x in outcomes
    if x in df.columns
]

treatments = [
    x for x in treatments
    if x in df.columns
]

covariates = [
    x for x in covariates
    if x in df.columns
]

# IMPUTE MISSING VALUES

all_cols = (
    outcomes +
    treatments +
    covariates
)

imputer = SimpleImputer(
    strategy="median"
)

df[all_cols] = imputer.fit_transform(
    df[all_cols]
)

# X MATRIX

X = df[covariates]

# STORE RESULTS

results = []

# CROSS-FITTING

kf = KFold(
    n_splits=5,
    shuffle=True,
    random_state=42
)

# LOOP THROUGH OUTCOMES

for outcome in outcomes:

    print("\n================================")
    print("Outcome:", outcome)
    print("================================")

    Y = df[outcome].values

    for treatment in treatments:

        print("\nTreatment:", treatment)

        T = df[treatment].values

        y_residuals = np.zeros(len(df))
        t_residuals = np.zeros(len(df))

        # CROSS-FITTING

        for train_idx, test_idx in kf.split(X):

            X_train = X.iloc[train_idx]
            X_test = X.iloc[test_idx]

            Y_train = Y[train_idx]
            Y_test = Y[test_idx]

            T_train = T[train_idx]
            T_test = T[test_idx]

            # MODEL Y ~ X

            model_y = XGBRegressor(

                n_estimators=300,
                max_depth=4,
                learning_rate=0.05,
                subsample=0.8,
                colsample_bytree=0.8,
                random_state=42,
                n_jobs=-1

            )

            model_y.fit(
                X_train,
                Y_train
            )

            y_hat = model_y.predict(
                X_test
            )

            y_residuals[test_idx] = (
                Y_test - y_hat
            )

            # MODEL T ~ X

            model_t = XGBRegressor(

                n_estimators=300,
                max_depth=4,
                learning_rate=0.05,
                subsample=0.8,
                colsample_bytree=0.8,
                random_state=42,
                n_jobs=-1

            )

            model_t.fit(
                X_train,
                T_train
            )

            t_hat = model_t.predict(
                X_test
            )

            t_residuals[test_idx] = (
                T_test - t_hat
            )

        theta = np.sum(
            t_residuals * y_residuals
        ) / np.sum(
            t_residuals ** 2
        )

        # STANDARD ERROR

        residual_final = (
            y_residuals -
            theta * t_residuals
        )

        sigma2 = np.mean(
            residual_final ** 2
        )

        se = np.sqrt(

            sigma2 /

            np.sum(
                t_residuals ** 2
            )

        )

        ci_low = theta - 1.96 * se
        ci_high = theta + 1.96 * se

        results.append({

            "Outcome": outcome,

            "Treatment": treatment,

            "Estimate": theta,

            "Std_Error": se,

            "CI_Lower": ci_low,

            "CI_Upper": ci_high

        })

# RESULTS DATAFRAME

results_df = pd.DataFrame(results)

# SAVE RESULTS

results_df.to_csv(

    "Manual_DML_Results.csv",

    index=False

)

print(results_df)

# PLOT RESULTS

for outcome in outcomes:

    plot_df = results_df[
        results_df["Outcome"] == outcome
    ].copy()

    plot_df = plot_df.sort_values(
        by="Estimate"
    )

    plt.figure(figsize=(10, 8))

    plt.errorbar(

        x=plot_df["Estimate"],

        y=plot_df["Treatment"],

        xerr=1.96 * plot_df["Std_Error"],

        fmt='o'

    )

    plt.axvline(
        0,
        linestyle='--'
    )

    plt.xlabel(
        "Orthogonalized Effect Estimate"
    )

    plt.ylabel(
        "Environmental Exposure"
    )

    plt.title(
        f"Manual DML Effects on {outcome.upper()}"
    )

    plt.tight_layout()

    plt.savefig(

        f"Manual_DML_{outcome.upper()}.png",

        dpi=300

    )

    plt.show()

# DONE

print("\nAnalysis Complete.")

print("\nGenerated Files:")

print("- Manual_DML_Results.csv")

for outcome in outcomes:

    print(
        f"- Manual_DML_{outcome.upper()}.png"
    )

# Environmental Exposures -> Epigenetic Aging

# INSTALL (RUN ONCE)

# IMPORT LIBRARIES

# CREATE DAG

G = nx.DiGraph()

# ADD NODES

# EXPOSURES

exposures = [

    "Precipitation",
    "Temperature",
    "PM2.5",
    "CO",
    "NO₂",
    "Ozone",
    "SO₂"

]

# HIV / IMMUNE PATHWAYS

immune_nodes = [

    "HIV Status",
    "Inflammation",
    "Immune Senescence",
    "Activated CD8",
    "CD4 Depletion"

]

# COVARIATES

covariates = [

    "Age",
    "Race",
    "BMI",
    "Smoking",
    "Education"

]

# OUTCOMES

outcomes = [

    "AAR",
    "EEAA",
    "PEAA",
    "GEAA",
    "DNAmTLadjAge"

]

# ADD TO GRAPH

for node in exposures:
    G.add_node(node, group="Exposure")

for node in immune_nodes:
    G.add_node(node, group="Immune")

for node in covariates:
    G.add_node(node, group="Covariate")

for node in outcomes:
    G.add_node(node, group="Outcome")

# ADD EDGES

# EXPOSURE -> IMMUNE

for exposure in exposures:

    G.add_edge(exposure, "Inflammation")
    G.add_edge(exposure, "Immune Senescence")

# HIV -> IMMUNE

G.add_edge("HIV Status", "Inflammation")
G.add_edge("HIV Status", "Immune Senescence")
G.add_edge("HIV Status", "Activated CD8")
G.add_edge("HIV Status", "CD4 Depletion")

# IMMUNE -> AGING

for outcome in outcomes:

    G.add_edge("Inflammation", outcome)
    G.add_edge("Immune Senescence", outcome)
    G.add_edge("Activated CD8", outcome)
    G.add_edge("CD4 Depletion", outcome)

# COVARIATES

for cov in covariates:

    for outcome in outcomes:
        G.add_edge(cov, outcome)

    for exposure in exposures:
        G.add_edge(cov, exposure)

# MANUAL NODE POSITIONS

pos = {

    # EXPOSURES
    "Precipitation": (-3, 4),
    "Temperature": (-3, 3),
    "PM2.5": (-3, 2),
    "CO": (-3, 1),
    "NO₂": (-3, 0),
    "Ozone": (-3, -1),
    "SO₂": (-3, -2),

    # HIV / IMMUNE
    "HIV Status": (-1, 2),
    "Inflammation": (1, 2),
    "Immune Senescence": (1, 0),
    "Activated CD8": (1, -2),
    "CD4 Depletion": (1, -4),

    # COVARIATES
    "Age": (-1, -6),
    "Race": (0, -6),
    "BMI": (1, -6),
    "Smoking": (2, -6),
    "Education": (3, -6),

    # OUTCOMES
    "AAR": (5, 4),
    "EEAA": (5, 2),
    "PEAA": (5, 0),
    "GEAA": (5, -2),
    "DNAmTLadjAge": (5, -4)
}

# NODE COLORS

node_colors = []

for node in G.nodes():

    group = G.nodes[node]["group"]

    if group == "Exposure":
        node_colors.append("skyblue")

    elif group == "Immune":
        node_colors.append("lightcoral")

    elif group == "Covariate":
        node_colors.append("lightgreen")

    elif group == "Outcome":
        node_colors.append("plum")

# CREATE FIGURE

plt.figure(figsize=(18, 12))

# DRAW NODES

nx.draw_networkx_nodes(

    G,
    pos,

    node_size=4500,

    node_color=node_colors,

    edgecolors="black",

    linewidths=1.5
)

# DRAW LABELS

nx.draw_networkx_labels(

    G,
    pos,

    font_size=11,

    font_weight="bold"
)

# DRAW EDGES

nx.draw_networkx_edges(

    G,
    pos,

    arrows=True,

    arrowstyle="-|>",

    arrowsize=25,

    width=2,

    connectionstyle="arc3,rad=0.05"
)

# TITLE

plt.title(

    "Directed Acyclic Graph (DAG): Environmental Exposures, Immune Dysregulation, and Epigenetic Aging",

    fontsize=20,

    fontweight="bold",

    pad=30
)

# LEGEND

legend_elements = [

    Patch(facecolor='skyblue', edgecolor='black', label='Environmental Exposures'),

    Patch(facecolor='lightcoral', edgecolor='black', label='HIV / Immune Pathways'),

    Patch(facecolor='lightgreen', edgecolor='black', label='Covariates'),

    Patch(facecolor='plum', edgecolor='black', label='Epigenetic Aging Outcomes')
]

plt.legend(

    handles=legend_elements,

    loc='upper left',

    fontsize=11,

    frameon=True
)

# REMOVE AXES

plt.axis("off")

# SAVE FIGURE

plt.savefig(

    "Publication_Quality_DAG.png",

    dpi=300,

    bbox_inches="tight"
)

plt.savefig(

    "Publication_Quality_DAG.pdf",

    bbox_inches="tight"
)

# SHOW FIGURE

plt.show()

# DONE

print("\nDAG figure generated successfully.")

print("\nFiles saved:")
print("- Publication_Quality_DAG.png")
print("- Publication_Quality_DAG.pdf")

# EXPORT DML RESULTS TO WORD (.docx)

# INSTALL (RUN ONCE)

# IMPORT LIBRARIES

# LOAD RESULTS

results_df = pd.read_csv(
    "Manual_DML_Results.csv"
)

# ROUND VALUES

results_df["Estimate"] = results_df["Estimate"].round(4)
results_df["Std_Error"] = results_df["Std_Error"].round(4)
results_df["CI_Lower"] = results_df["CI_Lower"].round(4)
results_df["CI_Upper"] = results_df["CI_Upper"].round(4)

# CREATE CI COLUMN

results_df["95% CI"] = (
    results_df["CI_Lower"].astype(str)
    + " to " +
    results_df["CI_Upper"].astype(str)
)

final_df = results_df[[

    "Outcome",
    "Treatment",
    "Estimate",
    "Std_Error",
    "95% CI"

]].copy()

# RENAME COLUMNS

final_df.columns = [

    "Outcome",
    "Environmental Exposure",
    "Effect Estimate",
    "Standard Error",
    "95% Confidence Interval"

]

# CREATE WORD DOCUMENT

doc = Document()

# LANDSCAPE PAGE

section = doc.sections[-1]

new_width, new_height = section.page_height, section.page_width

section.page_width = new_width
section.page_height = new_height

# TITLE

title = doc.add_heading(
    "Environmental Exposures and Epigenetic Aging: Manual DML Results",
    level=1
)

title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# DESCRIPTION

p = doc.add_paragraph()

p.add_run(
    "Table. Manual Double Machine Learning (DML) estimates for environmental exposures and epigenetic aging outcomes among Visit 4 participants. "
).bold = True

p.add_run(
    "Models adjusted for demographic, behavioral, and immune covariates. Viral load-related variables and HIV proxy interaction terms were excluded."
)

# CREATE TABLE

table = doc.add_table(
    rows=1,
    cols=len(final_df.columns)
)

table.style = 'Table Grid'

# HEADER ROW

hdr_cells = table.rows[0].cells

for i, col in enumerate(final_df.columns):

    hdr_cells[i].text = str(col)

    for paragraph in hdr_cells[i].paragraphs:

        for run in paragraph.runs:

            run.font.bold = True
            run.font.size = Pt(10)

# ADD DATA ROWS

for _, row in final_df.iterrows():

    row_cells = table.add_row().cells

    for i, value in enumerate(row):

        row_cells[i].text = str(value)

        for paragraph in row_cells[i].paragraphs:

            for run in paragraph.runs:

                run.font.size = Pt(9)

# AUTOFIT COLUMNS

for row in table.rows:

    for cell in row.cells:

        cell.width = Inches(1.5)

# FOOTNOTE

foot = doc.add_paragraph()

foot.add_run(
    "Abbreviations: AAR = Age Acceleration Residual; EEAA = Extrinsic Epigenetic Age Acceleration; PEAA = Phenotypic Epigenetic Age Acceleration; GEAA = GrimAge Epigenetic Age Acceleration; DNAmTLadjAge = DNA methylation telomere length adjusted age."
).italic = True

foot.style.font.size = Pt(8)

# SAVE DOCUMENT

doc.save(
    "Manual_DML_Results_Table.docx"
)

# DONE

print("\nWord document generated successfully.")

print("\nSaved file:")
print("- Manual_DML_Results_Table.docx")

---

# Output

This generates:

* `Manual_DML_Results_Table.docx`

with:

* publication-quality formatting
* landscape layout
* formatted confidence intervals
* bold headers
* readable font sizes
* manuscript-ready structure
# DOWHY CAUSAL ANALYSIS
# Environmental Exposures -> Epigenetic Aging
#
# Visit = 4 only
# Removes HIV proxy variables

# INSTALL (RUN ONCE)

# IMPORT LIBRARIES

# LOAD DATA

df = pd.read_csv("combined.csv")

# KEEP VISIT = 4 ONLY

df = df[df["visit"] == 4].copy()

print("\nVisit 4 Dataset Shape:")
print(df.shape)

# REMOVE HIV PROXY VARIABLES
# + INTERACTION TERMS

remove_keywords = [

    "cvl",
    "vload2",
    "esthivdate",
    "lg10cvl",
    "cvlfrhaart"

]

pattern = "|".join(remove_keywords)

df = df.loc[
    :,
    ~df.columns.str.contains(
        pattern,
        case=False,
        regex=True
    )
]

print("\nAfter Removing HIV Proxy Variables:")
print(df.shape)

# DEFINE OUTCOMES

outcomes = [

    "aar",
    "eeaa",
    "peaa",
    "geaa",
    "dnamtladjage"

]

# DEFINE TREATMENTS

treatments = [

    "Precipitation",
    "Temperature",
    "SPWPM2.5",
    "SBMPM2.5",
    "CO",
    "COAQI",
    "NO2",
    "NO2AQI",
    "Ozone",
    "OzoneAQI",
    "SO2",
    "SO2AQI"

]

# DEFINE COVARIATES

covariates = [

    "white",
    "educbas",
    "bmi",
    "cum_pkyear",

    "lncd3",
    "lncd4",
    "lncd8",

    "lnnaivecd4",
    "lnnaivecd8",

    "lnsencd4",
    "lnsencd8",

    "lnactcd4",
    "lnactcd8",

    "lnabs_cd42",
    "lnabs_cd82",

    "lnabsnaivecd4",
    "lnabsnaivecd8",

    "lnabssencd4",
    "lnabssencd8",

    "lnabsactcd4",
    "lnabsactcd8"

]

# KEEP EXISTING VARIABLES ONLY

outcomes = [
    x for x in outcomes
    if x in df.columns
]

treatments = [
    x for x in treatments
    if x in df.columns
]

covariates = [
    x for x in covariates
    if x in df.columns
]

# IMPUTE MISSING VALUES

all_cols = (
    outcomes +
    treatments +
    covariates
)

imputer = SimpleImputer(
    strategy="median"
)

df[all_cols] = imputer.fit_transform(
    df[all_cols]
)

# STORE RESULTS

results = []

# LOOP THROUGH OUTCOMES

for outcome in outcomes:

    print("\n================================================")
    print("Outcome:", outcome)
    print("================================================")

    for treatment in treatments:

        print("\nTreatment:", treatment)

        # BUILD DAG

        common_causes = covariates

        causal_graph = """

        digraph {

            """

        # Covariates -> Treatment

        for cov in common_causes:

            causal_graph += f"""
            {cov} -> {treatment};
            """

        # Covariates -> Outcome

        for cov in common_causes:

            causal_graph += f"""
            {cov} -> {outcome};
            """

        # Treatment -> Outcome

        causal_graph += f"""
        {treatment} -> {outcome};
        """

        causal_graph += """
        }
        """

        # CREATE CAUSAL MODEL

        model = CausalModel(

            data=df,

            treatment=treatment,

            outcome=outcome,

            graph=causal_graph

        )

        # IDENTIFY EFFECT

        identified_estimand = model.identify_effect()

        # ESTIMATE EFFECT

        estimate = model.estimate_effect(

            identified_estimand,

            method_name="backdoor.linear_regression"

        )

        # SAVE RESULTS

        results.append({

            "Outcome": outcome,

            "Treatment": treatment,

            "Estimate": estimate.value

        })

# RESULTS DATAFRAME

results_df = pd.DataFrame(results)

# SAVE RESULTS

results_df.to_csv(

    "DoWhy_Causal_Results.csv",

    index=False

)

print("\n================================================")
print("FINAL RESULTS")
print("================================================")

print(results_df)

# PLOT RESULTS

sns.set_style("whitegrid")

for outcome in outcomes:

    plot_df = results_df[
        results_df["Outcome"] == outcome
    ].copy()

    plot_df = plot_df.sort_values(
        by="Estimate"
    )

    plt.figure(figsize=(11, 8))

    bars = plt.barh(

        plot_df["Treatment"],

        plot_df["Estimate"]

    )

    plt.axvline(
        0,
        linestyle='--',
        linewidth=1.5
    )

    # ADD VALUE LABELS

    for bar in bars:

        width = bar.get_width()

        plt.text(

            width,

            bar.get_y() + bar.get_height()/2,

            f"{width:.3f}",

            va='center',

            fontsize=9

        )

    plt.xlabel(
        "Estimated Causal Effect",
        fontsize=12,
        fontweight='bold'
    )

    plt.ylabel(
        "Environmental Exposure",
        fontsize=12,
        fontweight='bold'
    )

    plt.title(

        f"DoWhy Estimated Effects on {outcome.upper()}",

        fontsize=16,

        fontweight='bold'

    )

    plt.tight_layout()

    plt.savefig(

        f"DoWhy_{outcome.upper()}_Effects.png",

        dpi=300,

        bbox_inches="tight"

    )

    plt.show()

# OVERALL HEATMAP

heatmap_df = results_df.pivot(

    index="Treatment",

    columns="Outcome",

    values="Estimate"

)

plt.figure(figsize=(10, 8))

sns.heatmap(

    heatmap_df,

    annot=True,

    fmt=".3f",

    cmap="coolwarm",

    center=0

)

plt.title(

    "DoWhy Estimated Environmental Effects Across Epigenetic Clocks",

    fontsize=16,

    fontweight='bold'

)

plt.tight_layout()

plt.savefig(

    "DoWhy_Heatmap.png",

    dpi=300,

    bbox_inches="tight"

)

plt.show()

# DONE

print("\nAnalysis Complete.")

print("\nGenerated Files:")

print("- DoWhy_Causal_Results.csv")

for outcome in outcomes:

    print(
        f"- DoWhy_{outcome.upper()}_Effects.png"
    )

print("- DoWhy_Heatmap.png")

# LOAD RESULTS

df = pd.read_csv(
    "Manual_DML_Results.csv"
)

# STYLE

sns.set_style("whitegrid")

# LOOP THROUGH OUTCOMES

outcomes = df["Outcome"].unique()

for outcome in outcomes:

    plot_df = df[
        df["Outcome"] == outcome
    ].copy()

    plot_df = plot_df.sort_values(
        by="Estimate"
    )

    # FIGURE

    plt.figure(figsize=(11, 8))

    # ERROR BARS

    plt.errorbar(

        x=plot_df["Estimate"],

        y=plot_df["Treatment"],

        xerr=1.96 * plot_df["Std_Error"],

        fmt='o',

        capsize=5,

        linewidth=2

    )

    # NULL LINE

    plt.axvline(

        0,

        linestyle='--',

        linewidth=2

    )

    # LABELS

    plt.xlabel(

        "Estimated Orthogonalized Causal Effect",

        fontsize=13,

        fontweight='bold'

    )

    plt.ylabel(

        "Environmental Exposure",

        fontsize=13,

        fontweight='bold'

    )

    # TITLE

    plt.title(

        f"Environmental Exposures and {outcome.upper()}",

        fontsize=18,

        fontweight='bold',

        pad=20

    )

    # VALUE LABELS

    for i, row in plot_df.iterrows():

        plt.text(

            row["Estimate"],

            row["Treatment"],

            f'{row["Estimate"]:.3f}',

            fontsize=9,

            va='center',

            ha='left'

        )

    # TIGHT LAYOUT

    plt.tight_layout()

    # SAVE FIGURE

    plt.savefig(

        f"Publication_Causal_{outcome.upper()}.png",

        dpi=300,

        bbox_inches='tight'

    )

    plt.show()

# OVERALL HEATMAP

heatmap_df = df.pivot(

    index="Treatment",

    columns="Outcome",

    values="Estimate"

)

plt.figure(figsize=(11, 8))

sns.heatmap(

    heatmap_df,

    annot=True,

    fmt=".3f",

    center=0,

    cmap="coolwarm",

    linewidths=0.5

)

plt.title(

    "Environmental Exposure Effects Across Epigenetic Clocks",

    fontsize=18,

    fontweight='bold',

    pad=20

)

plt.tight_layout()

plt.savefig(

    "Exposure_Outcome_Heatmap.png",

    dpi=300,

    bbox_inches='tight'

)

plt.show()

print("\nPublication-quality figures generated.")

# IMPORT LIBRARIES

# STYLE

plt.rcParams.update({

    "font.family": "sans-serif",

    "axes.linewidth": 1.2,

    "axes.labelsize": 13,

    "axes.titlesize": 16,

    "xtick.labelsize": 11,

    "ytick.labelsize": 11

})

sns.set_style("whitegrid")

# LOAD DATA

df = pd.read_csv("combined.csv")

# VISIT 4 ONLY

df = df[
    df["visit"] == 4
].copy()

# REMOVE HIV PROXY VARIABLES

remove_keywords = [

    "cvl",
    "vload2",
    "esthivdate",
    "lg10cvl",
    "cvlfrhaart"

]

pattern = "|".join(remove_keywords)

df = df.loc[
    :,
    ~df.columns.str.contains(
        pattern,
        case=False,
        regex=True
    )
]

# OUTCOMES

outcomes = [

    "aar",
    "eeaa",
    "peaa",
    "geaa",
    "dnamtladjage"

]

# EXPOSURES

exposures = [

    "Precipitation",
    "Temperature",
    "SPWPM2.5",
    "SBMPM2.5",
    "CO",
    "COAQI",
    "NO2",
    "NO2AQI",
    "Ozone",
    "OzoneAQI",
    "SO2",
    "SO2AQI"

]

# KEEP EXISTING VARIABLES

outcomes = [
    x for x in outcomes
    if x in df.columns
]

exposures = [
    x for x in exposures
    if x in df.columns
]

# IMPUTE MISSING

imputer = SimpleImputer(
    strategy="median"
)

df[
    outcomes + exposures
] = imputer.fit_transform(
    df[outcomes + exposures]
)

# NICE COLOR PALETTE

colors = [

    "#5B8FF9",
    "#61DDAA",
    "#65789B",
    "#F6BD16",

    "#7262FD",
    "#78D3F8",
    "#9661BC",
    "#F6903D",

    "#008685",
    "#F08BB4",
    "#C2C8D5",
    "#6DC8EC"

]

# LOOP THROUGH OUTCOMES

for outcome in outcomes:

    print("\n================================")
    print("Outcome:", outcome)
    print("================================")

    # FIGURE

    fig, axes = plt.subplots(

        3,
        4,

        figsize=(24, 18)

    )

    axes = axes.flatten()

    # LOOP THROUGH EXPOSURES

    for i, exposure in enumerate(exposures):

        ax = axes[i]

        x = df[exposure].values
        y = df[outcome].values

        # SCATTER

        ax.scatter(

            x,
            y,

            alpha=0.55,

            s=35,

            color=colors[i]

        )

        # POLYNOMIAL SMOOTHING

        sorted_idx = np.argsort(x)

        x_sorted = x[sorted_idx]
        y_sorted = y[sorted_idx]

        # cubic polynomial fit
        coeffs = np.polyfit(
            x_sorted,
            y_sorted,
            deg=3
        )

        poly = np.poly1d(coeffs)

        x_curve = np.linspace(
            x_sorted.min(),
            x_sorted.max(),
            300
        )

        y_curve = poly(x_curve)

        ax.plot(

            x_curve,
            y_curve,

            linewidth=3,

            color="#1F1F1F"

        )

        # DESCRIPTIVE STATS

        median = np.median(x)
        mean = np.mean(x)
        sd = np.std(x)

        xmin = np.min(x)
        xmax = np.max(x)

        corr = np.corrcoef(
            x,
            y
        )[0, 1]

        n = len(x)

        # STAT BOX

        stats_text = (

            f"N = {n}\n"

            f"Mean = {mean:.2f}\n"

            f"Median = {median:.2f}\n"

            f"SD = {sd:.2f}\n"

            f"Range = [{xmin:.2f}, {xmax:.2f}]\n"

            f"Correlation = {corr:.2f}"

        )

        ax.text(

            0.03,
            0.97,

            stats_text,

            transform=ax.transAxes,

            fontsize=9,

            verticalalignment='top',

            bbox=dict(

                boxstyle='round,pad=0.5',

                facecolor='white',

                edgecolor='gray',

                alpha=0.9

            )

        )

        # TITLES

        ax.set_title(

            exposure,

            fontsize=13,

            fontweight='bold'

        )

        ax.set_xlabel(

            exposure,

            fontsize=11

        )

        ax.set_ylabel(

            outcome.upper(),

            fontsize=11

        )

    # REMOVE EMPTY PANELS

    for j in range(
        len(exposures),
        len(axes)
    ):

        fig.delaxes(
            axes[j]
        )

    # MAIN TITLE

    plt.suptitle(

        f"Exposure–Response Relationships with {outcome.upper()}",

        fontsize=24,

        fontweight='bold',

        y=1.02

    )

    # TIGHT LAYOUT

    plt.tight_layout()

    # SAVE FIGURE

    plt.savefig(

        f"Exposure_Response_{outcome.upper()}.png",

        dpi=300,

        bbox_inches='tight'

    )

    plt.show()

# DONE

print("\nPublication-quality figures generated.")
# Plotting
for covariate_index in range(0, 25):
    covariate_name = data.columns[2 + covariate_index]  #covariate_name

    # Best fit line
    regressor = LinearRegression()
    regressor.fit(X_test[:, covariate_index].reshape(-1, 1), marginal_effects)
    best_fit_line = regressor.predict(X_test[:, covariate_index].reshape(-1, 1))

    # Plot the marginal effect of the treatment by the chosen covariate with the best fit line
    plt.figure(figsize=(8, 6))
    plt.scatter(X_test[:, covariate_index], marginal_effects, label='Marginal Effects')
    plt.plot(X_test[:, covariate_index], best_fit_line, color='red', label='Best Fit Line')
    plt.title(f'Marginal Effect of Treatment by {covariate_name}')
    plt.xlabel(covariate_name)
    plt.ylabel('Marginal Effect')
    plt.legend()
    plt.tight_layout()
    plt.savefig('XGBoost_dnamtladjage_' + str(covariate_name) + '.png', bbox_inches="tight", dpi=300)
    plt.show()

# IMPORT LIBRARIES

# LOAD RESULTS

results = pd.read_csv(
    "Manual_DML_Results.csv"
)

# LOAD DATA

df = pd.read_csv(
    "combined.csv"
)

df = df[
    df["visit"] == 4
].copy()

# REMOVE HIV PROXY VARIABLES

remove_keywords = [

    "cvl",
    "vload2",
    "esthivdate",
    "lg10cvl",
    "cvlfrhaart"

]

pattern = "|".join(remove_keywords)

df = df.loc[
    :,
    ~df.columns.str.contains(
        pattern,
        case=False,
        regex=True
    )
]

# STYLE

plt.rcParams.update({

    "font.family": "sans-serif",

    "axes.linewidth": 1.2,

    "axes.labelsize": 12,

    "axes.titlesize": 14

})

sns.set_style("whitegrid")

# SELECT TOP ASSOCIATIONS

results["abs_effect"] = np.abs(
    results["Estimate"]
)

top_results = results.sort_values(

    by="abs_effect",

    ascending=False

).head(6)

print(top_results)

# MULTI-PANEL FIGURE

fig, axes = plt.subplots(

    2,
    3,

    figsize=(18, 12)

)

axes = axes.flatten()

# COLORS

colors = [

    "#5B8FF9",
    "#61DDAA",
    "#65789B",
    "#F6BD16",
    "#7262FD",
    "#F6903D"

]

# LOOP THROUGH TOP EFFECTS

for i, (_, row) in enumerate(
    top_results.iterrows()
):

    outcome = row["Outcome"]
    exposure = row["Treatment"]

    ax = axes[i]

    x = df[exposure]
    y = df[outcome]

    # SCATTER

    ax.scatter(

        x,
        y,

        alpha=0.55,

        s=35,

        color=colors[i]

    )

    # POLYNOMIAL FIT

    sorted_idx = np.argsort(x)

    x_sorted = x.iloc[sorted_idx]
    y_sorted = y.iloc[sorted_idx]

    coeffs = np.polyfit(
        x_sorted,
        y_sorted,
        deg=3
    )

    poly = np.poly1d(coeffs)

    x_curve = np.linspace(
        x_sorted.min(),
        x_sorted.max(),
        300
    )

    y_curve = poly(x_curve)

    ax.plot(

        x_curve,
        y_curve,

        linewidth=3,

        color="black"

    )

    # STATS

    corr = np.corrcoef(
        x,
        y
    )[0, 1]

    estimate = row["Estimate"]

    stats_text = (

        f"Effect = {estimate:.3f}\n"

        f"Correlation = {corr:.2f}\n"

        f"N = {len(x)}"

    )

    ax.text(

        0.03,
        0.97,

        stats_text,

        transform=ax.transAxes,

        fontsize=9,

        verticalalignment='top',

        bbox=dict(

            boxstyle='round,pad=0.4',

            facecolor='white',

            alpha=0.9

        )

    )

    # TITLE

    ax.set_title(

        f"{exposure} → {outcome.upper()}",

        fontsize=13,

        fontweight='bold'

    )

    ax.set_xlabel(
        exposure
    )

    ax.set_ylabel(
        outcome.upper()
    )

# MAIN TITLE

plt.suptitle(

    "Top Environmental Exposure Associations with Epigenetic Aging",

    fontsize=22,

    fontweight='bold',

    y=1.02

)

# TIGHT LAYOUT

plt.tight_layout()

# SAVE

plt.savefig(

    "Top_Exposure_Response_Panels.png",

    dpi=600,

    bbox_inches='tight'

)

plt.show()

print("\nFigure saved:")
print("Top_Exposure_Response_Panels.png")

# IMPORT LIBRARIES

# STYLE

plt.rcParams.update({

    "font.family": "sans-serif",

    "axes.linewidth": 1.3,

    "axes.labelsize": 12,

    "axes.titlesize": 15,

    "xtick.labelsize": 10,

    "ytick.labelsize": 10

})

sns.set_style("whitegrid")

# LOAD DML RESULTS

results = pd.read_csv(
    "Manual_DML_Results.csv"
)

# LOAD DATA

df = pd.read_csv(
    "combined.csv"
)

# VISIT 4 ONLY

df = df[
    df["visit"] == 4
].copy()

# REMOVE HIV PROXY VARIABLES

remove_keywords = [

    "cvl",
    "vload2",
    "esthivdate",
    "lg10cvl",
    "cvlfrhaart"

]

pattern = "|".join(remove_keywords)

df = df.loc[
    :,
    ~df.columns.str.contains(
        pattern,
        case=False,
        regex=True
    )
]

# OUTCOMES

outcomes = [

    "aar",
    "eeaa",
    "peaa",
    "geaa",
    "dnamtladjage"

]

# IMPUTE MISSING

numeric_cols = df.select_dtypes(
    include=np.number
).columns

imputer = SimpleImputer(
    strategy="median"
)

df[numeric_cols] = imputer.fit_transform(
    df[numeric_cols]
)

# BEAUTIFUL COLOR PALETTE

colors = [

    "#3B82F6",   # blue
    "#10B981",   # emerald
    "#F59E0B",   # amber
    "#EF4444"    # red

]

# LOOP THROUGH OUTCOMES

for outcome in outcomes:

    print("\n================================")
    print("Outcome:", outcome)
    print("================================")

    # SELECT TOP 4 EXPOSURES

    sub = results[
        results["Outcome"] == outcome
    ].copy()

    sub["abs_est"] = np.abs(
        sub["Estimate"]
    )

    top4 = sub.sort_values(

        by="abs_est",

        ascending=False

    ).head(4)

    # FIGURE

    fig, axes = plt.subplots(

        2,
        2,

        figsize=(15, 12)

    )

    axes = axes.flatten()

    # LOOP THROUGH TOP EXPOSURES

    for i, (_, row) in enumerate(
        top4.iterrows()
    ):

        exposure = row["Treatment"]

        estimate = row["Estimate"]

        se = row["Std_Error"]

        ax = axes[i]

        x = df[exposure].values
        y = df[outcome].values

        # REMOVE EXTREME OUTLIERS

        lower = np.percentile(x, 1)
        upper = np.percentile(x, 99)

        mask = (x >= lower) & (x <= upper)

        x = x[mask]
        y = y[mask]

        # SCATTER

        ax.scatter(

            x,
            y,

            alpha=0.55,

            s=40,

            color=colors[i],

            edgecolor='none'

        )

        # POLYNOMIAL CURVE

        sorted_idx = np.argsort(x)

        x_sorted = x[sorted_idx]
        y_sorted = y[sorted_idx]

        coeffs = np.polyfit(

            x_sorted,
            y_sorted,
            deg=3

        )

        poly = np.poly1d(coeffs)

        x_curve = np.linspace(

            x_sorted.min(),
            x_sorted.max(),
            300

        )

        y_curve = poly(x_curve)

        ax.plot(

            x_curve,
            y_curve,

            linewidth=3.2,

            color="#111827"

        )

        # DESCRIPTIVE STATS

        mean = np.mean(x)
        median = np.median(x)
        sd = np.std(x)

        xmin = np.min(x)
        xmax = np.max(x)

        corr = np.corrcoef(
            x,
            y
        )[0, 1]

        n = len(x)

        # STAT BOX

        stats_text = (

            f"Effect = {estimate:.3f}\n"
            f"SE = {se:.3f}\n"
            f"N = {n}\n"
            f"Mean = {mean:.2f}\n"
            f"Median = {median:.2f}\n"
            f"SD = {sd:.2f}\n"
            f"Range = [{xmin:.2f}, {xmax:.2f}]\n"
            f"r = {corr:.2f}"

        )

        ax.text(

            0.03,
            0.97,

            stats_text,

            transform=ax.transAxes,

            fontsize=9,

            verticalalignment='top',

            bbox=dict(

                boxstyle='round,pad=0.5',

                facecolor='white',

                edgecolor='gray',

                alpha=0.92

            )

        )

        # TITLE

        ax.set_title(

            exposure,

            fontsize=15,

            fontweight='bold',

            pad=14

        )

        # AXIS LABELS

        ax.set_xlabel(

            exposure,

            fontsize=11,

            fontweight='bold'

        )

        ax.set_ylabel(

            outcome.upper(),

            fontsize=11,

            fontweight='bold'

        )

        # GRID

        ax.grid(

            alpha=0.3

        )

    # MAIN TITLE

    plt.suptitle(

        f"Top Environmental Associations with {outcome.upper()}",

        fontsize=24,

        fontweight='bold',

        y=1.02

    )

    # LAYOUT

    plt.tight_layout(
        pad=3.5
    )

    # SAVE FIGURE

    plt.savefig(

        f"Top4_Exposure_Response_{outcome.upper()}.png",

        dpi=600,

        bbox_inches='tight'

    )

    plt.show()

print("\n================================")
print("Publication-style figures generated.")
print("================================")
# Deep Learning + Causal Inference

# FIGURE SETUP

fig, ax = plt.subplots(figsize=(18, 16))

ax.set_xlim(0, 16)
ax.set_ylim(0, 18)

ax.axis('off')

# COLORS

BOX_COLOR = "#F8FAFC"
EDGE_COLOR = "#334155"
ARROW_COLOR = "#111827"

# DRAW BOX FUNCTION

def draw_box(
    x,
    y,
    text,
    width=3.0,
    height=1.1,
    fontsize=15
):

    box = FancyBboxPatch(

        (x, y),

        width,
        height,

        boxstyle="round,pad=0.05,rounding_size=0.12",

        linewidth=1.7,

        edgecolor=EDGE_COLOR,

        facecolor=BOX_COLOR

    )

    ax.add_patch(box)

    ax.text(

        x + width/2,
        y + height/2,

        text,

        ha='center',
        va='center',

        fontsize=fontsize,

        fontweight='semibold',

        color="#111827"

    )

# DRAW ARROW FUNCTION

def draw_arrow(x1, y1, x2, y2):

    ax.annotate(

        "",

        xy=(x2, y2),
        xytext=(x1, y1),

        arrowprops=dict(

            arrowstyle="->",

            linewidth=1.8,

            color=ARROW_COLOR,

            shrinkA=5,
            shrinkB=5

        )

    )

# MAIN TITLE

ax.text(

    8,
    17.1,

    "Deep Learning and Causal Inference Workflow",

    fontsize=30,

    fontweight='bold',

    ha='center',

    color="#111827"

)

# SECTION A

ax.text(

    0.6,
    15.5,

    "a. Data Collection",

    fontsize=22,

    fontweight='bold'

)

draw_box(
    1.0, 12.8,
    "Data Sources",
    width=2.5,
    fontsize=16
)

draw_box(
    5.0, 14.0,
    "Epigenetic Aging Data\nObtained from MACS",
    width=6.5,
    height=1.3,
    fontsize=15
)

draw_box(
    5.0, 12.1,
    "Climate Exposure Data\nObtained from NCEI",
    width=6.5,
    height=1.3,
    fontsize=15
)

draw_box(
    5.0, 10.0,
    "Air Pollution Exposure Data\nObtained from EPA and WashU",
    width=6.5,
    height=1.4,
    fontsize=15
)

draw_arrow(3.5, 13.3, 5.0, 14.6)
draw_arrow(3.5, 13.3, 5.0, 12.7)
draw_arrow(3.5, 13.3, 5.0, 10.8)

# SECTION B

ax.text(

    0.6,
    8.7,

    "b. Data Processing",

    fontsize=22,

    fontweight='bold'

)

draw_box(
    0.8, 6.0,
    "Removal of\nProxy Variables",
    width=2.8,
    height=1.4,
    fontsize=15
)

draw_box(
    4.2, 6.0,
    "Feature Engineering\nand Data Cleaning",
    width=3.2,
    height=1.4,
    fontsize=15
)

draw_box(
    8.2, 6.0,
    "Missing Data\nImputation",
    width=2.9,
    height=1.4,
    fontsize=15
)

draw_box(
    12.0, 6.0,
    "Normalization\nand Scaling",
    width=2.8,
    height=1.4,
    fontsize=15
)

draw_arrow(3.6, 6.7, 4.2, 6.7)
draw_arrow(7.4, 6.7, 8.2, 6.7)
draw_arrow(11.1, 6.7, 12.0, 6.7)

# SECTION C

ax.text(

    0.6,
    3.9,

    "c. Deep Learning and Exploratory Analysis",

    fontsize=22,

    fontweight='bold'

)

draw_box(
    2.2, 1.6,
    "Random Forest\nFeature Prioritization",
    width=3.5,
    height=1.5,
    fontsize=15
)

draw_box(
    6.4, 1.6,
    "K-Means\nClustering",
    width=3.0,
    height=1.5,
    fontsize=15
)

draw_box(
    10.2, 1.6,
    "Exposure Pattern\nDiscovery",
    width=3.5,
    height=1.5,
    fontsize=15
)

draw_arrow(5.7, 2.35, 6.4, 2.35)
draw_arrow(9.4, 2.35, 10.2, 2.35)

# SECTION D

ax.text(

    5.1,
    -0.2,

    "d. Causal Inference",

    fontsize=22,

    fontweight='bold'

)

draw_box(
    4.0, -2.4,
    "DAG-Guided\nCovariate Adjustment",
    width=3.6,
    height=1.5,
    fontsize=15
)

draw_box(
    8.6, -2.4,
    "XGBoost-Based\nDouble Machine Learning",
    width=4.2,
    height=1.5,
    fontsize=15
)

draw_arrow(7.6, -1.65, 8.6, -1.65)

# CONNECT SECTIONS

draw_arrow(8.0, 1.6, 8.0, -0.9)

plt.tight_layout()

# SAVE HIGH-RESOLUTION FIGURE

plt.savefig(

    "DeepLearning_CausalInference_Workflow.png",

    dpi=1000,

    bbox_inches='tight',

    facecolor='white'

)

plt.savefig(

    "DeepLearning_CausalInference_Workflow.pdf",

    bbox_inches='tight',

    facecolor='white'

)

plt.show()

print("\nPublication-quality workflow figure generated:")
print("- DeepLearning_CausalInference_Workflow.png")
print("- DeepLearning_CausalInference_Workflow.pdf")
# Deep Learning + Causal Inference

# FIGURE SETUP

fig, ax = plt.subplots(figsize=(18, 16))

# IMPORTANT:
# extend lower bound so causal boxes are visible
ax.set_xlim(0, 16)
ax.set_ylim(-4, 18)

ax.axis('off')

# JAMA-STYLE COLORS

BOX_COLOR = "#FFFFFF"
EDGE_COLOR = "#1F2937"
TITLE_COLOR = "#111827"
ARROW_COLOR = "#374151"

# FONT SETTINGS

TITLE_SIZE = 30
SECTION_SIZE = 22
BOX_SIZE = 15

# DRAW BOX FUNCTION

def draw_box(
    x,
    y,
    text,
    width=3.0,
    height=1.2,
    fontsize=15
):

    box = FancyBboxPatch(

        (x, y),

        width,
        height,

        boxstyle="round,pad=0.04,rounding_size=0.08",

        linewidth=1.8,

        edgecolor=EDGE_COLOR,

        facecolor=BOX_COLOR

    )

    ax.add_patch(box)

    ax.text(

        x + width/2,
        y + height/2,

        text,

        ha='center',
        va='center',

        fontsize=fontsize,

        fontweight='semibold',

        color=TITLE_COLOR

    )

# DRAW ARROW FUNCTION

def draw_arrow(x1, y1, x2, y2):

    ax.annotate(

        "",

        xy=(x2, y2),
        xytext=(x1, y1),

        arrowprops=dict(

            arrowstyle="->",

            linewidth=2.0,

            color=ARROW_COLOR,

            shrinkA=6,
            shrinkB=6

        )

    )

# MAIN TITLE

ax.text(

    8,
    17.0,

    "Deep Learning and Causal Inference Framework",

    fontsize=TITLE_SIZE,

    fontweight='bold',

    ha='center',

    color=TITLE_COLOR

)

# SECTION A

ax.text(

    0.6,
    15.0,

    "a. Data Collection",

    fontsize=SECTION_SIZE,

    fontweight='bold',

    color=TITLE_COLOR

)

draw_box(
    1.0, 12.2,
    "Data Sources",
    width=2.6,
    fontsize=BOX_SIZE
)

draw_box(
    5.0, 13.6,
    "Epigenetic Aging Data\nObtained from MACS",
    width=6.6,
    height=1.4,
    fontsize=BOX_SIZE
)

draw_box(
    5.0, 11.5,
    "Climate Exposure Data\nObtained from NCEI",
    width=6.6,
    height=1.4,
    fontsize=BOX_SIZE
)

draw_box(
    5.0, 9.2,
    "Air Pollution Exposure Data\nObtained from EPA and WashU",
    width=6.6,
    height=1.5,
    fontsize=BOX_SIZE
)

draw_arrow(3.6, 12.8, 5.0, 14.3)
draw_arrow(3.6, 12.8, 5.0, 12.2)
draw_arrow(3.6, 12.8, 5.0, 10.0)

# SECTION B

ax.text(

    0.6,
    7.8,

    "b. Data Processing",

    fontsize=SECTION_SIZE,

    fontweight='bold',

    color=TITLE_COLOR

)

draw_box(
    0.8, 5.0,
    "Removal of\nProxy Variables",
    width=2.9,
    height=1.5,
    fontsize=14
)

draw_box(
    4.3, 5.0,
    "Feature Engineering\nand Data Cleaning",
    width=3.3,
    height=1.5,
    fontsize=14
)

draw_box(
    8.4, 5.0,
    "Missing Data\nImputation",
    width=3.0,
    height=1.5,
    fontsize=14
)

draw_box(
    12.2, 5.0,
    "Normalization\nand Scaling",
    width=2.9,
    height=1.5,
    fontsize=14
)

draw_arrow(3.7, 5.75, 4.3, 5.75)
draw_arrow(7.6, 5.75, 8.4, 5.75)
draw_arrow(11.4, 5.75, 12.2, 5.75)

# SECTION C

ax.text(

    0.6,
    2.8,

    "c. Deep Learning and Exploratory Analysis",

    fontsize=SECTION_SIZE,

    fontweight='bold',

    color=TITLE_COLOR

)

draw_box(
    1.6, 0.2,
    "Random Forest\nFeature Prioritization",
    width=3.6,
    height=1.6,
    fontsize=14
)

draw_box(
    6.0, 0.2,
    "K-Means\nClustering",
    width=3.0,
    height=1.6,
    fontsize=14
)

draw_box(
    10.0, 0.2,
    "Exposure Pattern\nDiscovery",
    width=3.7,
    height=1.6,
    fontsize=14
)

draw_arrow(5.2, 1.0, 6.0, 1.0)
draw_arrow(9.0, 1.0, 10.0, 1.0)

# VERTICAL FLOW TO CAUSAL

draw_arrow(8.0, 0.2, 8.0, -1.0)

# SECTION D

ax.text(

    5.0,
    -1.8,

    "d. Causal Inference",

    fontsize=SECTION_SIZE,

    fontweight='bold',

    color=TITLE_COLOR

)

draw_box(
    3.2, -3.6,
    "DAG-Guided\nCovariate Adjustment",
    width=4.0,
    height=1.6,
    fontsize=14
)

draw_box(
    8.8, -3.6,
    "XGBoost-Based\nDouble Machine Learning",
    width=4.6,
    height=1.6,
    fontsize=14
)

draw_arrow(7.2, -2.8, 8.8, -2.8)

# LAYOUT

plt.tight_layout()

# SAVE

plt.savefig(

    "JAMA_Style_Framework.png",

    dpi=1200,

    bbox_inches='tight',

    facecolor='white'

)

plt.savefig(

    "JAMA_Style_Framework.pdf",

    bbox_inches='tight',

    facecolor='white'

)

plt.show()

print("\nJAMA-style workflow figure generated:")
print("- JAMA_Style_Framework.png")
print("- JAMA_Style_Framework.pdf")